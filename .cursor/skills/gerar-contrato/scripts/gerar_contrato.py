# -*- coding: utf-8 -*-
"""
Gerador de Contrato de Locação de Veículo (Lanza Locações).

Uso:
    python gerar_contrato.py dados.json

Caminhos em dados.json podem ser relativos à raiz do repositório (pasta worklanza).
"""
import copy
import datetime
import json
import os
import re
import shutil
import sys
from pathlib import Path

import docx
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[4]


def _lanza_paths_cfg():
    cfg = ROOT / "config" / "lanza_paths.json"
    if cfg.is_file():
        return json.loads(cfg.read_text(encoding="utf-8"))
    return {}


def _default_contratos_dir():
    p = _lanza_paths_cfg()
    return (p.get("contratosDir") or p.get("documentosRaiz") or str(ROOT / "contratos"))


# ----------------- valor por extenso (pt-BR / reais) -----------------
_U = ['zero','um','dois','três','quatro','cinco','seis','sete','oito','nove',
      'dez','onze','doze','treze','quatorze','quinze','dezesseis','dezessete','dezoito','dezenove']
_DEZ = ['','','vinte','trinta','quarenta','cinquenta','sessenta','setenta','oitenta','noventa']
_CEM = ['','cento','duzentos','trezentos','quatrocentos','quinhentos','seiscentos','setecentos','oitocentos','novecentos']

def _grupo(n):
    if n == 0: return ''
    if n == 100: return 'cem'
    out = []
    c, r = n // 100, n % 100
    if c: out.append(_CEM[c])
    if r:
        if r < 20:
            out.append(_U[r])
        else:
            d, u = r // 10, r % 10
            out.append(_DEZ[d] + (' e ' + _U[u] if u else ''))
    return ' e '.join(out)

def _inteiro_extenso(n):
    if n == 0: return 'zero'
    partes = []
    milhoes = n // 1000000
    milhares = (n % 1000000) // 1000
    resto = n % 1000
    if milhoes:
        partes.append((_grupo(milhoes) + (' milhões' if milhoes > 1 else ' milhão')))
    if milhares:
        partes.append('mil' if milhares == 1 else (_grupo(milhares) + ' mil'))
    if resto:
        partes.append(_grupo(resto))
    txt = partes[0]
    for i in range(1, len(partes)):
        ult = partes[i]
        usa_e = (i == len(partes) - 1) and (resto < 100 or resto % 100 == 0)
        txt += (' e ' if usa_e else ', ') + ult
    return txt

def valor_extenso(v):
    v = round(float(v) + 1e-9, 2)
    reais = int(v)
    cent = int(round((v - reais) * 100))
    txt = _inteiro_extenso(reais) + (' real' if reais == 1 else ' reais')
    if cent:
        txt += ' e ' + _inteiro_extenso(cent) + (' centavo' if cent == 1 else ' centavos')
    return txt

def brl(v):
    s = f'{float(v):,.2f}'
    return s.replace(',', 'X').replace('.', ',').replace('X', '.')

def cap(s):
    return s[:1].upper() + s[1:] if s else s

# ----------------- helpers docx -----------------
def set_paragraph_text(p, novo):
    for hl in p._element.findall(qn('w:hyperlink')):
        p._element.remove(hl)
    if not p.runs:
        p.add_run(novo)
        return
    p.runs[0].text = novo
    for r in p.runs[1:]:
        r.text = ''

def set_paragraph_rich(p, segments):
    segments = [(t, b) for (t, b) in segments if t]
    if not segments:
        return
    for hl in p._element.findall(qn('w:hyperlink')):
        p._element.remove(hl)
    if not p.runs:
        for t, b in segments:
            r = p.add_run(t)
            if b: r.bold = True
        return
    base = p.runs[0]
    base_rPr = base._element.find(qn('w:rPr'))
    base.text = segments[0][0]
    base.bold = True if segments[0][1] else None
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)
    for t, b in segments[1:]:
        r = p.add_run(t)
        if base_rPr is not None:
            r._element.insert(0, copy.deepcopy(base_rPr))
        r.bold = True if b else None

def B(t): return (t, True)
def N(t): return (t, False)

def bold_substring(segments, sub):
    if not sub:
        return segments
    out = []
    for t, b in segments:
        if b or sub not in t:
            out.append((t, b)); continue
        i = t.index(sub)
        if t[:i]: out.append((t[:i], False))
        out.append((sub, True))
        if t[i+len(sub):]: out.append((t[i+len(sub):], False))
    return out

def rich_from_pattern(text, pat, troca):
    segs, last = [], 0
    for m in pat.finditer(text):
        if m.start() > last:
            segs.append(N(text[last:m.start()]))
        novo = troca(m)
        segs.append((novo, novo != m.group(0)))
        last = m.end()
    if last < len(text):
        segs.append(N(text[last:]))
    return segs

def fipe_url_mes_atual(url):
    if not url:
        return url
    hoje = datetime.date.today()
    return re.sub(r'\d{1,2}-\d{4}', f"{hoje.month}-{hoje.year}", url, count=1)

def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)

_MESES = ['janeiro','fevereiro','março','abril','maio','junho','julho','agosto',
          'setembro','outubro','novembro','dezembro']
_UF_NOME = {
    'SC':'Santa Catarina','SP':'São Paulo','RS':'Rio Grande do Sul','PR':'Paraná',
    'RJ':'Rio de Janeiro','MG':'Minas Gerais','BA':'Bahia','PE':'Pernambuco',
    'CE':'Ceará','GO':'Goiás','DF':'Distrito Federal','ES':'Espírito Santo',
}

def parse_data(s):
    return datetime.datetime.strptime(s, '%d/%m/%Y')

# ----------------- geração -----------------
def gerar(dados):
    doc = docx.Document(dados['template'])
    cli = dados['cliente']
    end = cli.get('endereco', {})
    veic = dados['veiculo']
    prazo = dados['prazo']
    val = dados['valores']
    diaria = float(val.get('diaria', 120.0))
    semana = float(val['semana'])
    caucao = float(val['caucao'])
    tres_diarias = diaria * 3
    categoria = dados.get('cnhCategoria', 'B')

    hora = prazo.get('hora', '18:00')
    dini = parse_data(prazo['inicio'])
    dfim = dini + datetime.timedelta(days=int(prazo['dias']))
    inicio_str = f"{dini.strftime('%d/%m/%Y')} as {hora}"
    fim_str = f"{dfim.strftime('%d/%m/%Y')} as {hora}"

    estado_cli = end.get('estado') or _UF_NOME.get((end.get('uf') or '').upper(), end.get('uf', ''))
    log = end.get('logradouro', '')
    if end.get('numero'): log += f", {end['numero']}"
    if end.get('complemento'): log += f", {end['complemento']}"

    mm = veic.get('marcaModelo', '')
    fm = (veic.get('fipeModelo') or '').strip()
    if fm:
        mm += f" ({fm})"

    fipe_url = fipe_url_mes_atual(veic.get('fipe', '') or '')

    for p in doc.paragraphs:
        if p.text.strip().startswith('LOCAT') and 'CPF' in p.text:
            set_paragraph_rich(p, [
                N("LOCATÁRIO(a): "), B(cli['nome']),
                N(", inscrito no CPF sob o n° "), B(cli['cpf']),
                N(", residente e domiciliado na "), B(log),
                N(", bairro "), B(end.get('bairro', '')),
                N(", cidade "), B(end.get('cidade', '')),
                N(", estado "), B(estado_cli),
                N(", CEP "), B(end.get('cep', '')), N("."),
            ])
            break

    seg_v = [N("1.1 O presente contrato tem como OBJETO a locação do automóvel de placa: "),
             B(veic['placa']), N(", marca/modelo "), B(mm)]
    if veic.get('chassi'):  seg_v += [N(", Chassi "), B(veic['chassi'])]
    if veic.get('renavam'): seg_v += [N(", RENAVAM "), B(veic['renavam'])]
    seg_v += [N(", ano/modelo "), B(veic['anoModelo']), N(", cor "), B(veic['cor']), N(".")]

    veh_pars  = [p for p in doc.paragraphs if p.text.strip().startswith('1.1 O presente')]
    fipe_pars = [p for p in doc.paragraphs if p.text.strip().startswith('1.1.1')]
    if veh_pars:
        set_paragraph_rich(veh_pars[0], seg_v)
        for p in veh_pars[1:]: delete_paragraph(p)
    if fipe_pars:
        set_paragraph_rich(fipe_pars[0], [N("1.1.1 Informações tabela Fipe: "), B(fipe_url)])
        for p in fipe_pars[1:]: delete_paragraph(p)

    for p in doc.paragraphs:
        if p.text.strip().startswith('1.2 A presente'):
            set_paragraph_rich(p, [
                N("1.2 A presente locação terá o lapso temporal de validade de "),
                B(f"{int(prazo['dias'])} dias"),
                N(", podendo ser renovado por vontade das partes, iniciando no dia "),
                B(inicio_str), N(" e terminando no dia "), B(fim_str),
                N(", com tolerância de no máximo 1 hora. Data na qual o automóvel deverá ser "
                  "devolvido no estado em que foi locado, sem avarias."),
            ])
            break

    catpat = re.compile(r'(categoria\s*["“])([^"”]*)(["”])')
    for p in doc.paragraphs:
        if p.text.strip().startswith('2.1.') and 'categoria' in p.text:
            t = p.text
            m = catpat.search(t)
            if m:
                set_paragraph_rich(p, [
                    N(t[:m.start()]), N(m.group(1)), B(categoria), N(m.group(3)), N(t[m.end():]),
                ])
            break

    alvos = {
        650.00: (semana, valor_extenso(semana)),
        120.00: (diaria, valor_extenso(diaria)),
        1500.00: (caucao, valor_extenso(caucao)),
        360.00: (tres_diarias, valor_extenso(tres_diarias)),
    }
    pat = re.compile(r'R\$\s*([\d.]+,\d{2})\s*\(([^)]*)\)')
    def troca(m):
        num = float(m.group(1).replace('.', '').replace(',', '.'))
        for ref, (nv, ext) in alvos.items():
            if abs(num - ref) < 0.005:
                return f"R$ {brl(nv)} ({cap(ext)})"
        return m.group(0)
    dia_pag = dados.get('diaPagamento')
    for p in doc.paragraphs:
        if 'R$' in p.text and pat.search(p.text):
            t = p.text
            is_32 = t.strip().startswith('3.2 O pagamento')
            if is_32 and dia_pag:
                t = t.replace('todas as segundas-feiras', dia_pag)
            novo = pat.sub(troca, t)
            if novo != p.text:
                segs = rich_from_pattern(t, pat, troca)
                if is_32 and dia_pag:
                    segs = bold_substring(segs, dia_pag)
                set_paragraph_rich(p, segs)

    assin = dados.get('assinatura', {})
    data_assin = assin.get('data', 'auto')
    if data_assin == 'auto':
        data_assin = f"{dini.day:02d} de {_MESES[dini.month-1]} de {dini.year}"
    cidade_a = assin.get('cidade', 'Tubarão')
    estado_a = assin.get('estado', 'Santa Catarina')
    for p in doc.paragraphs:
        if re.search(r'\d{1,2} de [A-Za-zçãéíóúâ]+ de \d{4}', p.text) and 'lapso' not in p.text:
            set_paragraph_rich(p, [N(f"{cidade_a}, {estado_a}, "), B(data_assin), N(".")])
            break

    for p in doc.paragraphs:
        if 'JOSE FELIPE BARRETO' in p.text.upper() and 'LOCAT' not in p.text.upper():
            m = re.search(r'(\t+|\s{2,})(JOSE FELIPE BARRETO RODRIGUES)', p.text, flags=re.IGNORECASE)
            if m:
                set_paragraph_rich(p, [B(cli['nome']), N(p.text[m.start():])])
            break

    nome_cli = cli['nome'].title()
    base_dir = dados.get('contratosDir') or _default_contratos_dir()
    pasta = os.path.join(base_dir, f"{dini.strftime('%d.%m.%Y')} - {nome_cli}")
    os.makedirs(pasta, exist_ok=True)
    nome_arq = f"Contrato - {nome_cli}"
    saida_docx = os.path.join(pasta, nome_arq + '.docx')
    saida_pdf  = os.path.join(pasta, nome_arq + '.pdf')

    doc.save(saida_docx)

    pdf_ok = False
    try:
        import win32com.client, pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        wb = word.Documents.Open(os.path.abspath(saida_docx))
        wb.SaveAs(os.path.abspath(saida_pdf), FileFormat=17)
        wb.Close(False)
        word.Quit()
        pdf_ok = True
    except Exception as e:
        print(f'[aviso] PDF nao gerado: {e}')

    cnh_dest = None
    cnh_src = dados.get('cnhArquivo')
    if cnh_src and os.path.exists(cnh_src):
        cnh_dest = os.path.join(pasta, 'CNH.pdf')
        shutil.copy(cnh_src, cnh_dest)

    return {"pasta": pasta, "docx": saida_docx,
            "pdf": saida_pdf if pdf_ok else None, "cnh": cnh_dest}

def _abs_repo(p):
    if not p:
        return p
    pp = Path(p)
    return str(pp) if pp.is_absolute() else str((ROOT / pp).resolve())

if __name__ == '__main__':
    with open(sys.argv[1], encoding='utf-8') as f:
        dados = json.load(f)
    for k in ('template', 'contratosDir', 'cnhArquivo'):
        if dados.get(k):
            dados[k] = _abs_repo(dados[k])
    r = gerar(dados)
    print(f'Pasta -> {r["pasta"]}')
    print(f'Word  -> {r["docx"]}')
    if r["pdf"]:
        print(f'PDF   -> {r["pdf"]}')
    if r["cnh"]:
        print(f'CNH   -> {r["cnh"]}')
    else:
        print('[aviso] CNH.pdf nao copiada (informe "cnhArquivo" no dados.json)')
